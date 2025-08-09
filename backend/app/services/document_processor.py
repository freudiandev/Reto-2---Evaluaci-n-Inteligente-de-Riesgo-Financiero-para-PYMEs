import pandas as pd
import openpyxl
from PyPDF2 import PdfReader
from docx import Document
import json
import re
from typing import Dict, List, Optional, Tuple
import os

class FinancialDocumentProcessor:
    """Procesador de documentos financieros para extraer estados financieros"""
    
    def __init__(self):
        # Patrones para identificar conceptos financieros
        self.financial_patterns = {
            'activos_corrientes': [
                'activos corrientes', 'activo corriente', 'current assets',
                'efectivo', 'bancos', 'cuentas por cobrar', 'inventarios'
            ],
            'activos_no_corrientes': [
                'activos no corrientes', 'activo no corriente', 'non-current assets',
                'propiedad planta equipo', 'activos fijos', 'inmuebles'
            ],
            'pasivos_corrientes': [
                'pasivos corrientes', 'pasivo corriente', 'current liabilities',
                'cuentas por pagar', 'deudas corto plazo'
            ],
            'pasivos_no_corrientes': [
                'pasivos no corrientes', 'pasivo no corriente', 'non-current liabilities',
                'deudas largo plazo', 'préstamos bancarios'
            ],
            'patrimonio': [
                'patrimonio', 'equity', 'capital social', 'utilidades retenidas'
            ],
            'ingresos': [
                'ingresos', 'ventas', 'revenue', 'facturación', 'ingresos operacionales'
            ],
            'costos': [
                'costo de ventas', 'cost of goods sold', 'costos operacionales'
            ],
            'gastos_operacionales': [
                'gastos operacionales', 'gastos administrativos', 'gastos de ventas'
            ],
            'gastos_financieros': [
                'gastos financieros', 'intereses', 'financial expenses'
            ]
        }

    def process_file(self, file_path: str) -> Dict:
        """Procesar archivo y extraer información financiera"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        processors = {
            '.xlsx': self._process_excel,
            '.xls': self._process_excel,
            '.pdf': self._process_pdf,
            '.docx': self._process_word,
            '.txt': self._process_text
        }
        
        if file_extension in processors:
            try:
                return processors[file_extension](file_path)
            except Exception as e:
                return {'error': f'Error procesando archivo: {str(e)}'}
        else:
            return {'error': f'Tipo de archivo no soportado: {file_extension}'}

    def _process_excel(self, file_path: str) -> Dict:
        """Procesar archivo Excel"""
        try:
            # Leer todas las hojas
            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                sheets_data[sheet_name] = df
            
            # Buscar datos financieros en las hojas
            financial_data = self._extract_financial_data_from_sheets(sheets_data)
            
            return {
                'type': 'excel',
                'sheets': list(sheets_data.keys()),
                'financial_data': financial_data,
                'raw_data': self._convert_dataframes_to_dict(sheets_data)
            }
            
        except Exception as e:
            return {'error': f'Error procesando Excel: {str(e)}'}

    def _process_pdf(self, file_path: str) -> Dict:
        """Procesar archivo PDF"""
        try:
            reader = PdfReader(file_path)
            text_content = ""
            
            for page in reader.pages:
                text_content += page.extract_text() + "\n"
            
            # Extraer datos financieros del texto
            financial_data = self._extract_financial_data_from_text(text_content)
            
            return {
                'type': 'pdf',
                'pages': len(reader.pages),
                'text_content': text_content,
                'financial_data': financial_data
            }
            
        except Exception as e:
            return {'error': f'Error procesando PDF: {str(e)}'}

    def _process_word(self, file_path: str) -> Dict:
        """Procesar archivo Word"""
        try:
            doc = Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extraer datos financieros del texto
            financial_data = self._extract_financial_data_from_text(text_content)
            
            return {
                'type': 'word',
                'paragraphs': len(doc.paragraphs),
                'text_content': text_content,
                'financial_data': financial_data
            }
            
        except Exception as e:
            return {'error': f'Error procesando Word: {str(e)}'}

    def _process_text(self, file_path: str) -> Dict:
        """Procesar archivo de texto"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            # Extraer datos financieros del texto
            financial_data = self._extract_financial_data_from_text(text_content)
            
            return {
                'type': 'text',
                'text_content': text_content,
                'financial_data': financial_data
            }
            
        except Exception as e:
            return {'error': f'Error procesando texto: {str(e)}'}

    def _extract_financial_data_from_sheets(self, sheets_data: Dict[str, pd.DataFrame]) -> Dict:
        """Extraer datos financieros de hojas de Excel"""
        financial_data = {
            'current_assets': 0,
            'non_current_assets': 0,
            'current_liabilities': 0,
            'non_current_liabilities': 0,
            'equity': 0,
            'total_revenue': 0,
            'cost_of_goods_sold': 0,
            'operating_expenses': 0,
            'financial_expenses': 0,
            'net_income': 0
        }
        
        for sheet_name, df in sheets_data.items():
            # Convertir DataFrame a texto para buscar patrones
            sheet_text = df.to_string().lower()
            
            # Buscar y extraer valores financieros
            for concept, patterns in self.financial_patterns.items():
                value = self._find_financial_value(df, patterns)
                if value and value != 0:
                    # Mapear conceptos a campos del modelo
                    field_mapping = {
                        'activos_corrientes': 'current_assets',
                        'activos_no_corrientes': 'non_current_assets',
                        'pasivos_corrientes': 'current_liabilities',
                        'pasivos_no_corrientes': 'non_current_liabilities',
                        'patrimonio': 'equity',
                        'ingresos': 'total_revenue',
                        'costos': 'cost_of_goods_sold',
                        'gastos_operacionales': 'operating_expenses',
                        'gastos_financieros': 'financial_expenses'
                    }
                    
                    if concept in field_mapping:
                        financial_data[field_mapping[concept]] = value
        
        # Calcular campos derivados
        financial_data['total_assets'] = financial_data['current_assets'] + financial_data['non_current_assets']
        financial_data['total_liabilities'] = financial_data['current_liabilities'] + financial_data['non_current_liabilities']
        financial_data['gross_profit'] = financial_data['total_revenue'] - financial_data['cost_of_goods_sold']
        financial_data['operating_income'] = financial_data['gross_profit'] - financial_data['operating_expenses']
        financial_data['net_income'] = financial_data['operating_income'] - financial_data['financial_expenses']
        
        return financial_data

    def _extract_financial_data_from_text(self, text: str) -> Dict:
        """Extraer datos financieros de texto"""
        text_lower = text.lower()
        financial_data = {}
        
        # Buscar patrones numéricos cerca de conceptos financieros
        for concept, patterns in self.financial_patterns.items():
            for pattern in patterns:
                # Buscar el patrón y extraer números cercanos
                matches = re.finditer(re.escape(pattern), text_lower)
                for match in matches:
                    # Buscar números en un rango de 100 caracteres después del patrón
                    start = match.end()
                    end = min(start + 100, len(text))
                    text_segment = text[start:end]
                    
                    # Extraer números (incluyendo decimales y comas)
                    numbers = re.findall(r'[\d,]+\.?\d*', text_segment)
                    if numbers:
                        try:
                            value = float(numbers[0].replace(',', ''))
                            financial_data[concept] = value
                            break
                        except ValueError:
                            continue
        
        return financial_data

    def _find_financial_value(self, df: pd.DataFrame, patterns: List[str]) -> Optional[float]:
        """Buscar valor financiero en DataFrame basado en patrones"""
        for pattern in patterns:
            # Buscar en todas las celdas del DataFrame
            for col in df.columns:
                for idx, row in df.iterrows():
                    cell_value = str(row[col]).lower()
                    if pattern in cell_value:
                        # Buscar números en la misma fila o columnas adyacentes
                        for search_col in df.columns:
                            try:
                                value = df.loc[idx, search_col]
                                if pd.notna(value) and isinstance(value, (int, float)):
                                    return float(value)
                                elif pd.notna(value):
                                    # Intentar convertir string a número
                                    clean_value = str(value).replace(',', '').replace('$', '').replace('(', '-').replace(')', '')
                                    if re.match(r'^-?\d+\.?\d*$', clean_value):
                                        return float(clean_value)
                            except (ValueError, TypeError):
                                continue
        return None

    def _convert_dataframes_to_dict(self, sheets_data: Dict[str, pd.DataFrame]) -> Dict:
        """Convertir DataFrames a diccionarios serializables"""
        result = {}
        for sheet_name, df in sheets_data.items():
            # Convertir DataFrame a diccionario, manejando valores NaN
            result[sheet_name] = df.fillna('').to_dict(orient='records')
        return result

    def validate_financial_data(self, financial_data: Dict) -> Tuple[bool, List[str]]:
        """Validar coherencia de datos financieros"""
        errors = []
        
        # Validar balance general
        total_assets = financial_data.get('current_assets', 0) + financial_data.get('non_current_assets', 0)
        total_liabilities = financial_data.get('current_liabilities', 0) + financial_data.get('non_current_liabilities', 0)
        equity = financial_data.get('equity', 0)
        
        if abs(total_assets - (total_liabilities + equity)) > 0.01:
            errors.append("El balance general no cuadra: Activos ≠ Pasivos + Patrimonio")
        
        # Validar que los valores sean positivos donde corresponde
        positive_fields = ['current_assets', 'non_current_assets', 'total_revenue']
        for field in positive_fields:
            if financial_data.get(field, 0) < 0:
                errors.append(f"{field} no puede ser negativo")
        
        # Validar coherencia de estado de resultados
        revenue = financial_data.get('total_revenue', 0)
        costs = financial_data.get('cost_of_goods_sold', 0)
        if costs > revenue and revenue > 0:
            errors.append("Los costos de ventas no pueden ser mayores a los ingresos")
        
        return len(errors) == 0, errors

class SuperciasDataExtractor:
    """Extractor de datos del portal de la Superintendencia de Compañías"""
    
    def __init__(self):
        self.base_url = "https://www.supercias.gob.ec/portalscvs/"
        
    def search_company_by_ruc(self, ruc: str) -> Dict:
        """Buscar empresa por RUC en el portal de Supercias"""
        # Por ahora simularemos la búsqueda ya que requiere integración real
        return self._simulate_supercias_data(ruc)
    
    def _simulate_supercias_data(self, ruc: str) -> Dict:
        """Simular datos de Supercias para desarrollo"""
        return {
            'ruc': ruc,
            'company_name': 'EMPRESA EJEMPLO S.A.',
            'legal_status': 'ACTIVA',
            'activity': 'COMERCIO AL POR MAYOR Y AL POR MENOR',
            'province': 'PICHINCHA',
            'canton': 'QUITO',
            'registration_date': '2018-03-15',
            'capital': 50000.0,
            'shareholders': 2,
            'legal_representative': 'JUAN PÉREZ',
            'address': 'AV. AMAZONAS N123 Y COLÓN',
            'phone': '02-2234567',
            'email': 'info@empresaejemplo.com',
            'has_financial_statements': True,
            'last_financial_statement_year': 2023,
            'simulated': True
        }
